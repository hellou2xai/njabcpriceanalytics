#!/usr/bin/env python
"""Generate the step-by-step video walkthrough scripts (Word .docx).

A production script for recording a SET of short user-guidance videos of CELR
Retail Pricing Intelligence - one video per numbered left-nav group:
  Video 1 Overview (Dashboard, Alerts)
  Video 2 Find Deals (Catalog, New Items, Combos, RIP Products)
  Video 3 Favorites
  Video 4 To-Do & Notes
  Video 5 Ordering (Cart, Lists, Orders)
  Video 6 Setup (Configuration)
Each video is self-contained: what it covers, a target length, a spoken intro
and outro, and its own scenes. Each scene gives the on-screen steps (how to do
it), the exact voiceover, which screen/screenshot to capture, and how it helps.

Run: python scripts/make_video_script.py
Output: docs/CELR_Video_Guide_Script.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "CELR_Video_Guide_Script.docx"

BLUE = RGBColor(0x1F, 0x4E, 0x8C)
SLATE = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x55, 0x63, 0x70)

# Each scene: (title, steps[list], screenshot, voiceover, why)
# Each video: dict(num, title, covers, length, intro, outro, scenes[])

VIDEOS = [
    {
        "num": 1,
        "title": "Overview",
        "covers": "Signing in, then the Overview group in the menu: Dashboard and Alerts.",
        "length": "2 to 3 minutes",
        "intro": "Welcome to CELR Retail Pricing Intelligence. CELR reads every New Jersey wholesaler price book each month and turns hundreds of pages into one searchable tool, with all the discount and rebate maths already done. In this first short video we will sign in and tour your home base - the Dashboard and your Alerts - so you know where the day's biggest deals and warnings show up the moment you log in.",
        "outro": "That is your home base. The Dashboard shows the shape of the month at a glance, and Alerts tells you what needs attention today. In the next video we go hunting for deals in the Catalog.",
        "scenes": [
            ("Open the app and sign in",
             ["Go to nj.celr.ai.",
              "On the sign-in screen, enter your email and password and click Sign in.",
              "First-time users: click Create an account, fill name, email, phone, password, tick the Terms box, then activate via the email link."],
             "Screen-record the sign-in screen, then the dashboard loading after login.",
             "First, head to nj.celr.ai and sign in with your email and password. If this is your first time, click Create an account, add your details, accept the terms, and activate from the email we send you.",
             "Gets the viewer into the app so every later step has context."),
            ("Dashboard",
             ["Click Dashboard (top of the left menu).",
              "Use the distributor buttons at the top to filter the whole page to one wholesaler.",
              "Click any of the six metric cards (Total Items, Active Discounts, Clearance, Price Drops, Price Increases, Active RIPs) to jump to that list.",
              "Scroll to the deal tiles (Time-Sensitive Deals, Biggest Price Drops, Top Discounts) and click one to open the full list."],
             "Full dashboard (cards + deal tiles). Then a click into one tile to show the drill-down.",
             "This is your daily starting point. The cards across the top are the shape of this month's price book at a glance: how many items, how many active discounts and the total dollars on the table, clearance items, and what went up or down versus last month. Each card is clickable. Below, the opportunity tiles do the hunting for you - the biggest price drops, the deals about to expire, the largest savings. Click any tile to open the full, sortable list.",
             "Shows the user where to get the day's highlights without searching, and that everything is one click from the headline number."),
            ("Alerts",
             ["Click Alerts (the bell; a red badge shows unread).",
              "Read the two groups: Opportunities (don't miss) and Watch-outs (avoid a mistake).",
              "Click a tile to jump to the screen with the detail.",
              "Click Mark all read to clear the badge."],
             "The Alerts page showing the grouped tiles; hover the unread badge in the nav.",
             "Alerts is an automatic digest - it builds itself, no buttons to press, and refreshes overnight. It is split into Opportunities, like deals ending this week, new rebates, and a favourite hitting your target price; and Watch-outs, like an order that is a couple of cases short of a bigger rebate, or an item that is cheaper at another distributor. Click any tile to act on it.",
             "Tells the user the app proactively flags what matters, so they never miss an expiring deal or a better price."),
        ],
    },
    {
        "num": 2,
        "title": "Find Deals",
        "covers": "The Find Deals group: Catalog, New Items, Combos, and RIP Products - plus the product details popup and the right-click menu you use everywhere.",
        "length": "5 to 7 minutes (the longest video; record it in parts)",
        "intro": "This video is all about finding deals. We will work through the four screens under Find Deals - Catalog, New Items, Combos, and RIP Products - and the product details popup and right-click menu you will use constantly. It is the longest video, so feel free to watch it in sections.",
        "outro": "That is how you find any deal in CELR - by product, by what is new, by bundle, or by rebate, with the real cost always in front of you. Next we will look at Favorites, your personal watchlist.",
        "scenes": [
            ("Catalog - open it and search",
             ["Click Catalog in the left menu.",
              "Type a product name or a barcode in the search box.",
              "Note the result count beside the box updating as you type."],
             "Catalog landing: the search box, the result count, and the first rows with product images.",
             "The Catalog is every product from every distributor in one searchable place. Type a name or a barcode in the search box - and every word counts, so 'bds sangria' finds BDS Pink Sangria even though those words are not next to each other. The number beside the box tells you how many products match.",
             "Shows the user the single fastest way to find any product out of tens of thousands."),
            ("Catalog - the distributor buttons (top right)",
             ["Point to the row of buttons in the top right: All Distributors, Allied, Fedway, Highgrade, Opici, Peerless.",
              "Click one distributor to filter the whole page to that wholesaler.",
              "Click All Distributors to bring everyone back."],
             "Zoom on the top-right distributor buttons; click Allied, show the list and counts change, then All Distributors.",
             "Across the top right are your distributors - All Distributors, Allied, Fedway, Highgrade, Opici, Peerless. Click any one to filter the entire catalogue, and all its counts, to just that supplier. Click All Distributors to see everyone together again. Use this when you only want to shop one wholesaler today.",
             "Lets the user focus on a single supplier in one click instead of scanning everyone."),
            ("Catalog - the filter panel (left)",
             ["Click Show / Hide Filters to toggle the panel.",
              "Deals: tick Has RIP offer, No RIP, Has discount, No discount, or In combo (bundle products).",
              "Distributors: multi-select wholesalers (with counts).",
              "Brand: search or tick a brand (verified names); Show all to expand.",
              "Price Range (Case): enter Min/Max and click Go.",
              "Category: Wine, Spirits, Beer, RTD, and so on. Size: every pack size, smallest to largest."],
             "The full filter panel; tick a couple of filters (e.g. Has discount + a brand) and show the list narrowing with counts.",
             "On the left, the filter panel narrows the catalogue fast. Deals lets you show only products with a rebate, a discount, or that are part of a combo. Distributors is a multi-select version of the top buttons. Brand uses verified product names - search or pick. Then there is a case price range, product Category, and Size, listed smallest to largest. Every option shows a count so you know how many products you will get before you click.",
             "Turns a giant list into exactly the slice the user wants - by deal, brand, price, category, or size."),
            ("Catalog - reading the columns",
             ["Walk left to right across one row's columns.",
              "PRODUCT: image, name, barcode. A blue 'Multiple distributors' tag means the same product is carried by several wholesalers; for wine, a 'Vintage' tag shows the year (one barcode can cover several vintages, each priced separately). A green 'In combo' tag means it is in a bundle (click it to open the bundle in a popup).",
              "CASE / BTL = list price; TIER = how many quantity tiers exist; SAVE = best saving; EFFECTIVE = what you actually pay after the best discount and rebate.",
              "ROI / GP% = return and margin; BETTER PRICE = whether it is cheaper this month or next."],
             "Slow pan across the column headers, then highlight the Multiple-distributors tag, In combo, the EFFECTIVE column, and a Better Price badge.",
             "Each row reads left to right. First the product - its image, name, and barcode. A blue Multiple distributors tag means the same product is carried by more than one wholesaler, so you can compare and buy from whoever is cheaper; for wine a Vintage tag shows the year, since one barcode can cover several vintages and each is priced separately; a green In combo tag means it is part of a bundle you can open in a popup. Then the list Case and Bottle price, how many tiers it has, the best saving, and - the number that really matters - the Effective price, what you pay after the best discount and rebate. ROI and GP percent show your return and margin.",
             "Teaches the user to read the real cost at a glance, not just the sticker price, and to read the data tags and bundles correctly."),
            ("Catalog - the Better Price column",
             ["Point to the BETTER PRICE column and its badge (SAME, This Month, or Next Month).",
              "Hover the badge to see this month's vs next month's effective price."],
             "Zoom on the Better Price column with a SAME badge, then a row showing This Month / Next Month.",
             "The Better Price badge answers one question: should I buy now or wait? SAME means the price is steady; This Month means it gets more expensive next month, so buy now; Next Month means it drops, so it may pay to wait. Hover for the exact this-month versus next-month prices.",
             "Tells the user the best time to buy, turning a price book into timing advice."),
            ("Catalog - discount & RIP tiers",
             ["Look under a deal product at the indented tier rows.",
              "Blue DISC rows are price-list discounts; green RIP rows are rebates.",
              "Read each tier: 'Buy N cs = $X', the saving, the effective price, and the ROI at that quantity.",
              "Increase the quantity and watch a tier light up green as you reach it."],
             "Close-up of the stacked tier rows under one product (Buy 5 cs = $15, Buy 20 cs = $75, ...), with the green highlight as qty hits a tier.",
             "Under a product with deals, the indented rows are its quantity tiers. Blue DISC rows are discounts off the price list; green RIP rows are rebates you earn back. Each one reads 'buy this many cases for this much off', with the resulting effective price and the ROI. As you raise the quantity, the tier you have unlocked lights up green. This is where you decide how many to buy to hit the best price.",
             "Shows the user exactly how much more to buy to unlock a deeper saving - the core money decision."),
            ("Catalog - add to cart",
             ["Set the Qty steppers (cases and/or bottles) on a row.",
              "Click the + (Add to cart) button in the Order column.",
              "Watch the cart count badge in the top-right increase.",
              "Click the star to favourite, or click the row to open the full product details popup."],
             "Set a quantity, click +, and show the top-right cart badge incrementing.",
             "When a product is worth buying, set the cases and bottles with the steppers and click the plus in the Order column. It drops straight into your cart, top-right, where it will be grouped by sales rep ready to send. Click the star to save it to Favorites, or click the row for the deep product view with price history and every tier.",
             "Teaches the core daily action - going from a deal to an order line in one click."),
            ("Product details popup - the Price Breakdown",
             ["Click any product row to open the details popup.",
              "Read the top: product image, category path, description, and specs (size, ABV, region).",
              "Note the four headline numbers: Case Cost, Bottle Cost, Best Discount, and Effective.",
              "Follow the Price Breakdown chart left to right: List, minus Discount, minus RIP, equals You pay.",
              "Scroll down for the full discount/RIP tier tables, the all-editions history, the price chart, and your notes."],
             "The product popup focused on the waterfall chart (green List bar, red Discount and RIP reductions, green You-pay bar).",
             "Click any product to open its full picture. The chart in the middle is the easiest way to understand the deal - think of your money flowing downhill. The tall green bar on the left is the list price, sixty dollars a case here. The two red bars are what comes off: the discount takes away fifteen dollars, and the rebate another five. The green bar on the right is what you actually pay - forty dollars. So in one glance: list price, minus the discount, minus the rebate, equals your real cost. Below the chart you get every tier spelled out, the price history, and a place for your own notes.",
             "Demystifies how the effective price is calculated, so the user trusts the number and sees the deal at a glance."),
            ("The right-click menu (works on any product, on any page)",
             ["On ANY product row - Catalog, New Items, Combos, RIP Products, Favorites, Lists, even search results - right-click it, or click the three-dot (...) button, to open the quick-actions menu.",
              "View Product: open the details popup.",
              "Search the web: look up street/retail prices and product info online.",
              "Add to Cart: drop it into your cart, ready to send to the rep.",
              "Add to Favorites: star it for your watchlist.",
              "Add to List: add it to one of your named lists (or create a new one).",
              "Add to To-Do: capture a dated task about this product.",
              "Copy Code: copy the barcode to your clipboard."],
             "Right-click a product row (do it once on the Catalog and once on another page) to show the same menu opening; pan the action list.",
             "Here is a shortcut that works everywhere. Right-click any product row - on any page - or click its three-dot button, and you get the same quick-actions menu. From it you can open the full details, search the web for retail prices, add the product to your cart, star it to Favorites, drop it onto a list, set a To-Do about it, or copy its barcode. It is the fastest way to act on a product without leaving the page you are on.",
             "One consistent menu on every screen means the user never has to navigate away to act on a product."),
            ("New Items",
             ["Click New Items.",
              "Use the month buttons at the top to show items introduced in a given month.",
              "Read the Introduced column; everything else works like the Catalog."],
             "New Items page with the month buttons and the Introduced column.",
             "New Items is the catalogue filtered to products that are genuinely new this month - their barcode was not in last month's book. It is the fastest way to spot what just launched and decide whether to bring it in.",
             "Helps the user stay first-to-shelf on new products without scanning the whole catalogue."),
            ("Combos - overview and the summary cards",
             ["Click Combos.",
              "Read the four summary cards: total Combos, Avg Savings, Max Savings, Avg Discount."],
             "The Combos page header with the four green summary cards.",
             "Combos are multi-product bundles sold as one pack. The four cards at the top frame the whole picture: how many bundles are on offer, the average dollar saving, the single biggest saving available, and the average percentage off. It is an instant read on how rich the bundle deals are this month.",
             "Gives the user the lay of the land on bundles before they dig in."),
            ("Combos - filters (Distributor, Min savings, Validity)",
             ["Combo description search at the top of the filter panel.",
              "Distributor: All, Allied, Fedway, Highgrade, Opici, Peerless.",
              "Min Savings: Any, $5+, $10+, $25+, $50+, $100+.",
              "Validity: Any, Valid this month, Valid next month, Valid both months."],
             "Zoom on the filter panel; click 'Valid this month', then 'Valid next month', and show the list change.",
             "Filter the bundles down on the left. Pick a distributor, set a minimum dollar saving so only the meaningful bundles show, and use Validity to choose when the deal runs - valid this month, valid next month, or valid in both. That last one is gold: bundles you can count on across the changeover.",
             "Lets the user zero in on bundles worth their time and plan around when each one is live."),
            ("Combos - reading the columns and the Outlook",
             ["Walk the columns: Combo (name + item count), Distributor, Items, Combo Price, Regular Value, % Off, Savings, Next Mo. Save, Valid Through.",
              "Focus on the OUTLOOK badge: Stable, 'Buy now - ends this month', or 'New next month'.",
              "Note the amber warning icon on % Off when a figure looks unusual."],
             "Slow pan across the columns; then highlight the Outlook column with its three badge types and the amber % Off warning.",
             "Each row compares the bundle price against the regular value of buying the items separately, with the dollar saving and percent off. Next Mo. Save shows whether the deal carries into next month. The Outlook badge is the headline: Stable means it continues, 'Buy now - ends this month' means grab it before it expires, and 'New next month' means it starts soon. If a percent-off shows an amber warning, the distributor's figures look unusual - verify before relying on it.",
             "Tells the user the real saving and, crucially, the urgency - what to grab now versus what is coming."),
            ("Combos - the bundle breakdown",
             ["Click any combo row to open its breakdown.",
              "Review each item in the pack: regular price, combo price, and the saving on each.",
              "Use Add bundle to Cart to drop the whole pack into your cart.",
              "From the Catalog, the green 'In combo' tag opens the same breakdown in a popup window."],
             "A combo row expanding/opening into its item-by-item breakdown.",
             "Click any bundle to see exactly what is inside: every product, its regular price, its combo price, and the saving it contributes - so you can see which items carry the deal. From here you can add the whole bundle to your cart in one go.",
             "Shows the user precisely where a bundle's value comes from, so they buy with confidence."),
            ("RIP Products (rebates)",
             ["Click RIP Products.",
              "Compare this month and next month side by side for each rebate program.",
              "Filter by distributor, incentive type, minimum saving or GP%, and 'new next month'."],
             "RIP Products with the this-month vs next-month columns and the Better badge.",
             "RIP Products is built around rebate programs. For every product with a rebate you see this month next to next month - the tier, the saving, the effective price - plus a badge telling you whether to buy now or wait. Use it to time your volume buys before a program changes.",
             "Helps the user chase the rebates worth chasing and time them right."),
        ],
    },
    {
        "num": 3,
        "title": "Favorites",
        "covers": "The Favorites screen (your watchlist).",
        "length": "1 to 2 minutes",
        "intro": "A quick one. Favorites is your personal watchlist - the handful of products you want to keep an eye on. Let me show you how to build it and how it tells you when to act.",
        "outro": "So Favorites keeps your shortlist in one place and pings you when a price hits your target. Next, turning the things you find into tracked tasks with To-Do and Notes.",
        "scenes": [
            ("Favorites",
             ["Star a product anywhere (catalogue, search, popup) to add it here.",
              "Click Favorites in the left menu.",
              "Set a Target price; add an inline note; group by category.",
              "Use the row + or quantities to send favourites to the cart."],
             "Favorites page with a target price set and the buy-signal/trend column.",
             "Anything you star lands in Favorites. It is your shortlist for watching prices: set a target price and the app alerts you when the market reaches it, see a buy signal and a trend arrow on each item, and add a quick note. When you are ready, push favourites straight into the cart.",
             "Lets the user track the handful of products they care about and get told when to act."),
        ],
    },
    {
        "num": 4,
        "title": "To-Do & Notes",
        "covers": "The My Work tools: the To-Do board and Notes.",
        "length": "2 to 3 minutes",
        "intro": "This video covers your two organising tools: the To-Do board, so a follow-up never slips, and Notes, where everything you jot down lives in one place.",
        "outro": "That keeps your follow-ups dated and your notes together. Now let's turn all of this into orders you send to your reps.",
        "scenes": [
            ("To-Do board",
             ["Add a task two ways: right-click any product anywhere and choose Add to To-Do (set the task, an optional note, and a due date), or use the + on a column / 'Add one here' for a standalone task.",
              "Read the columns left to right: Past (overdue), This week, Next week, In 2 weeks, and 3+ weeks / Later. Each shows its date range and a count.",
              "On a card: the task title, the due date, a pencil to edit, a circle to mark done, and a bin to delete; click the product name to open it.",
              "Drag a card to another column to reschedule it (the due date updates).",
              "Overdue tasks sit in red in the Past column until done; finished tasks collect in a Done list at the bottom."],
             "The To-Do board: the Past/overdue column in red plus the four weekly buckets; drag a card from one week to another.",
             "To-Do is a simple board so a follow-up never slips. Add a task two ways: right-click any product and choose Add to To-Do - the product and the page you were on are saved with it so you have the context later - or add a standalone task with the plus on a column. Tasks are sorted into buckets by when they are due: a red Past column for anything overdue, then This week, Next week, In two weeks, and three-plus weeks or later, each with its dates and a count. On a card you can edit it, tick it done, or delete it, and click the product to jump to it. Best of all, just drag a card to another week to reschedule - the due date moves with it. Overdue items stay red in Past until you deal with them.",
             "Makes sure a 'come back to this' never gets lost - every follow-up is dated, sorted by urgency, and reschedulable with a drag."),
            ("Notes",
             ["Click Notes.",
              "Write a sticky note (title, text, colour) and click Add note.",
              "Below the stickies, browse the single feed of every note you wrote on products, favourites, orders, and order lines; filter or search it."],
             "Notes page: the sticky-note grid on top and the combined notes feed below.",
             "Notes gives you sticky notes for quick reminders, and below them a single feed of every note you have written anywhere in the app - on a product, an order, a favourite - so nothing gets lost. Each note links back to where it came from.",
             "Keeps all the user's scattered notes in one searchable place."),
        ],
    },
    {
        "num": 5,
        "title": "Ordering: Cart, Lists & Orders",
        "covers": "The whole ordering flow - the Cart (top-right), your saved Lists, and submitted Orders.",
        "length": "4 to 5 minutes",
        "intro": "This is where it all comes together: turning the deals you found into orders you send to your reps. We will walk the full flow, then look at the Cart, your reusable Lists, and your submitted Orders. One thing first - if you have not added your sales reps yet, watch the Setup video, because the cart emails orders to the rep on file.",
        "outro": "And that is the whole loop: browse, add to cart, send to your reps, and revise later if you need to. The last short video covers the one-time Setup that powers all of this.",
        "scenes": [
            ("The full flow: browse - add to cart - save for later - send to rep",
             ["Browse and find products anywhere (Catalog, New Items, Combos, RIP Products, Favorites, or a saved List).",
              "Add what you want: set the quantity and click the + on a row, or right-click > Add to Cart. The cart count in the top-right ticks up.",
              "Open the Cart (top-right icon). Items are automatically grouped by the sales rep who covers that distributor.",
              "Fine-tune each line: set cases/bottles to reach a better tier, add a line note, and add a header note for the rep.",
              "Not buying an item yet? Click Save for later to park it below; click Move to cart to bring it back anytime.",
              "Check each group's sales rep, then click Send All Orders to Reps - one purchase order per rep is created and emailed.",
              "Find it later under Orders to review, revise (reopen + re-submit), cancel, or re-share."],
             "End-to-end sequence: browse + click +, the cart badge filling, the rep-grouped cart with tiers + notes + a Save-for-later item, then the Send button and the success message.",
             "Here is the whole journey, start to finish. You browse - the catalogue, new items, a bundle, your favourites, or a list you saved. When something is worth buying, set the quantity and hit the plus, or right-click and Add to Cart; you will see the cart count climb in the top corner. Open the cart and everything is already sorted by sales rep. Now fine-tune: nudge a quantity to unlock a better tier, jot a note on a line or for the whole order. If you are not ready on something, Save it for later - it drops into a holding area and you can pull it back whenever. When the rep on each group looks right, click Send All Orders to Reps: the app turns each group into a purchase order and emails it to that rep. And it is not locked away - find it under Orders any time to revise and re-send, cancel, or re-share. That is the full loop: browse, cart, send.",
             "Ties every feature together into the one workflow that matters - turning browsing into a sent, rep-ready order, with a parking lot for maybes."),
            ("The Cart",
             ["Add products with the + on a row or right-click > Add to Cart.",
              "Open the Cart from the icon in the top-right corner (it shows a live count).",
              "Items group by sales rep; pick or change the rep per distributor group.",
              "Adjust quantities to hit a tier; add a line note and a header note; use Save for later.",
              "Click Send All Orders to Reps to create and email one order per rep."],
             "The Cart: a distributor group with the rep dropdown, deal info and tiers per line, the notes, and the Send button.",
             "The cart is where you turn deals into an order. Add products from anywhere with the plus button or right-click. In the cart, items are grouped by the sales rep who covers that distributor - pick the rep if there is more than one. Each line shows the same deal info as the catalogue and its tiers, so you can bump a quantity to hit a better price right here. Add a note to a line or to the whole order, park anything in Save for later, and when you are ready, one click - Send All Orders to Reps - creates and emails a purchase order to each rep.",
             "This is the heart of the workflow: it shows the user how to go from browsing to a sent, rep-ready order in one place, hitting better tiers before they commit."),
            ("Lists",
             ["Right-click any product > Add to List (pick a list or create a new one).",
              "Click Lists; create as many named lists as you want with the + button.",
              "Tick items and click Move to cart (they stay in the list) or Delete selected."],
             "The Lists page: My lists on the left, a selected list with checkboxes and the Move-to-cart / Delete buttons.",
             "Lists are reusable, named collections - a seasonal reset, a promo, a regular reorder. Build a list once by right-clicking products and choosing Add to List. Later, tick the items you want and move them to the cart in one go; the list stays intact so you can reuse it next time.",
             "Saves the user from rebuilding the same buy every cycle - plan once, reorder forever."),
            ("Orders (submitted)",
             ["Click Orders to see the orders you have sent.",
              "Open an order to review lines, totals, RIP rebates, and GP%.",
              "Reopen to edit and re-submit as a new revision; cancel; or preview/re-share the PO PDF."],
             "Orders list, then an open order with totals and the Submit/Reopen/Preview-PDF controls.",
             "Orders holds the purchase orders you have sent. Open one to review every line, the invoice total, the rebate you will earn back, and your margin. Need a change? Reopen it, edit, and submit again as a new revision - your rep gets the updated PO. You can also cancel an order or re-share the PDF.",
             "Reassures the user that a sent order is still editable and traceable, with the rep always kept in sync."),
        ],
    },
    {
        "num": 6,
        "title": "Setup: Configuration",
        "covers": "The one-time Setup under Configuration: Sales Reps, Divisions, and Stores.",
        "length": "1 to 2 minutes",
        "intro": "One last short video on setup. Do this once and the ordering flow just works. Under Configuration you tell CELR about your sales reps, your divisions, and your stores.",
        "outro": "That is the one-time setup done, and that is the whole app. Find the deal, drop it in the cart, send it to your rep - with the maths done for you the whole way. For a written walkthrough of any screen, open the How To Guide in the menu, and use the Feedback button any time to reach us. Happy buying.",
        "scenes": [
            ("Configuration",
             ["Click Configuration in the left menu.",
              "Sales Reps tab: add a rep with name, distributor (required), optional division, email, and phone, then Add.",
              "Divisions tab: add your grouping labels (each tied to a distributor).",
              "Stores tab: add or edit your store locations (the address is looked up for you)."],
             "Configuration with the three tabs; show adding a sales rep, and the rep table with the email/phone.",
             "Before you send orders, set up your master data under Configuration. Add each sales rep with their distributor and, importantly, their email - that is where a sent order's purchase order goes. Add your divisions and your store locations too. The cart uses these reps to group your order and to know who to email, so this one-time setup makes the whole ordering flow click into place.",
             "Explains the one-time setup that powers rep grouping and the auto-emailed POs - without it, Send to reps cannot work."),
        ],
    },
]


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CELR Retail Pricing Intelligence")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("User-Guidance Video Series - Production Scripts")
    rs.font.size = Pt(13); rs.font.color.rgb = MUTED
    doc.add_paragraph()

    # How to use
    h = doc.add_heading("How to use these scripts", level=1)
    for run in h.runs:
        run.font.color.rgb = BLUE
    for line in [
        "This document is a SET of short videos - one per menu section. Record and publish each one separately so a user can watch just the part they need.",
        "Record at 1080p (1920x1080), browser maximised, light theme, sidebar expanded.",
        "Each video has an Intro line and an Outro line to read, then numbered scenes. Each scene gives four things: How to do it (the clicks to perform on screen), Screenshot / capture (what to show or grab), Voiceover (read this aloud, calm and clear, ~150 words per minute), and Why it helps (the benefit to mention or imply).",
        "Keep each scene to 20-45 seconds. Use a gentle zoom or a highlight box on the element being described.",
        "Demo account data: use a store with a few favourites and at least one sales rep set up (with an email) so the Cart and order flow show real grouping.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()

    # Contents
    hc = doc.add_heading("The videos", level=1)
    for run in hc.runs:
        run.font.color.rgb = BLUE
    for v in VIDEOS:
        p = doc.add_paragraph(style="List Bullet")
        rb = p.add_run(f"Video {v['num']} - {v['title']}: ")
        rb.bold = True
        p.add_run(f"{v['covers']} ({v['length']}).")

    for v in VIDEOS:
        # Each video starts on a new page
        doc.add_page_break()

        vh = doc.add_heading(f"Video {v['num']} - {v['title']}", level=1)
        for run in vh.runs:
            run.font.color.rgb = BLUE

        meta = doc.add_paragraph()
        meta.add_run("Covers: ").bold = True
        meta.add_run(v["covers"])
        ln = doc.add_paragraph()
        ln.add_run("Target length: ").bold = True
        ln.add_run(v["length"])

        intro = doc.add_paragraph()
        intro.add_run("Intro (voiceover): ").bold = True
        ir = intro.add_run("“" + v["intro"] + "”")
        ir.italic = True
        doc.add_paragraph()

        for i, (title, steps, screenshot, voiceover, why) in enumerate(v["scenes"], start=1):
            sc = doc.add_heading(f"Scene {v['num']}.{i}: {title}", level=2)
            for run in sc.runs:
                run.font.color.rgb = SLATE

            doc.add_paragraph().add_run("How to do it (on screen):").bold = True
            for s in steps:
                doc.add_paragraph(s, style="List Number")

            cap = doc.add_paragraph()
            cap.add_run("Screenshot / capture: ").bold = True
            cap.add_run(screenshot)

            vo = doc.add_paragraph()
            vo.add_run("Voiceover: ").bold = True
            vr = vo.add_run("“" + voiceover + "”")
            vr.italic = True

            wy = doc.add_paragraph()
            wy.add_run("Why it helps: ").bold = True
            wy.add_run(why)

            doc.add_paragraph()

        outro = doc.add_paragraph()
        outro.add_run("Outro (voiceover): ").bold = True
        orn = outro.add_run("“" + v["outro"] + "”")
        orn.italic = True

    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
